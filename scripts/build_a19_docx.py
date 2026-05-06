"""Generate the A-19 District IT Chairperson 6-page Word (.docx) deck.

Mirrors the PDF deck: page 1 is a cover with 4 logos in the header and
a 2x2 grid of 4 governor photo placeholders; pages 2-6 carry topics 2-6
each with a 2x3 image gallery (6 placeholders).

Run:  python3 scripts/build_a19_docx.py
Output: docs/A19_IT_CHAIRPERSON_DECK.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

LIONS_BLUE = RGBColor(0x00, 0x3F, 0x87)
LIONS_GOLD = RGBColor(0xFF, 0xC7, 0x2C)
DARK_GREY = RGBColor(0x2B, 0x2B, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUT = Path(__file__).resolve().parent.parent / "docs" / "A19_IT_CHAIRPERSON_DECK.docx"

COVER_LOGOS = ["LCI", "MD 3232", "DIST 3232F1", "A-19"]

GOVERNORS = [
    ("Int'l President", "Lions International"),
    ("Council Chair", "Multiple District 3232"),
    ("District Governor", "District 3232F1"),
    ("IPDG / VDG", "District 3232F1"),
]

PAGES = [
    {
        "kind": "cover",
        "page_label": "PAGE 1 · TOPIC 1",
        "title": "Roles & Responsibilities",
        "subtitle": "District Information Technology Chairperson · 3232F1",
        "lead": (
            "Lions International IT Chairperson — Job Duties, Responsibilities & "
            "Benefits. Branding Lionism for impact at District and Club levels."
        ),
        "bullets": [
            "Digital leader for the District Cabinet & 100+ clubs",
            "Governance: digital policy, data privacy, brand standards",
            "Systems: MyLion · MyLCI · Lion Portal · District website",
            "Club support: onboarding, training, helpdesk for every club",
            "Security: RLS, backups, HMAC webhooks, DPDP Act 2023",
            "Reporting: monthly MMR · dashboards · Annual Activity Report",
        ],
    },
    {
        "kind": "content",
        "page_label": "PAGE 2 · TOPIC 2",
        "title": "Knowledge of Lionism, Website, Portal & LLC",
        "subtitle": "Lion Portal · Website · Lions Learning Center · Digital Ecosystem",
        "bullets": [
            "Lionism — 'We Serve' · 5 Global Causes · 1.4M+ Lions",
            "Lion Portal — SSO, officer reporting, dues, MMR",
            "Website — Next.js, SEO, WCAG 2.1 AA, mobile-first",
            "Lions Learning Center — courses for every officer",
            "MyLion / MyLCI — service activities & membership",
            "Brand: Lions Blue #003F87 · Gold #FFC72C · We Serve voice",
        ],
        "gallery": [
            "Lion Portal", "Lionism", "LLC Course",
            "District Website", "MyLion App", "Brand Kit",
        ],
    },
    {
        "kind": "content",
        "page_label": "PAGE 3 · TOPIC 3",
        "title": "Preparation & Implementation of District IT Goals",
        "subtitle": "Plan · Execute · Monitor",
        "bullets": [
            "Annual IT roadmap aligned to DG's theme",
            "Quarterly milestones with club-level targets",
            "100% MMR compliance across every club",
            "Every club onboarded with website + social presence",
            "Real-time dashboards for the Cabinet",
            "Zero data incidents — security & backups audited",
        ],
        "gallery": [
            "Cabinet Meet", "IT Goals Chart", "Roadmap",
            "Dashboard", "Quarterly Review", "Audit Report",
        ],
    },
    {
        "kind": "content",
        "page_label": "PAGE 4 · TOPIC 4",
        "title": "E-Directory · Website · E-Club House · Multimedia",
        "subtitle": "Club · Zone · Region · District Level",
        "bullets": [
            "E-Directory — searchable, role-aware, mobile-ready",
            "Website — district + 50+ club micro-sites, unified brand",
            "E-Club House — virtual meets, RSVPs, QR check-ins",
            "Multimedia — photo & video archive at every level",
            "Brand-safe templates: posters, certificates, social cards",
            "Self-service publishing: clubs go live in minutes",
        ],
        "gallery": [
            "E-Directory", "Club Website", "E-Club House",
            "Photo Archive", "Video Reel", "Template Pack",
        ],
    },
    {
        "kind": "content",
        "page_label": "PAGE 5 · TOPIC 5",
        "title": "Social Media Promotion at Digital Platforms",
        "subtitle": "Reach · Engagement · Branding of Lionism",
        "bullets": [
            "Platforms: Facebook · Instagram · LinkedIn · YouTube · X",
            "10,000+ combined followers — +180% YoY growth",
            "Weekly campaign: #WeServeWednesday",
            "Reels & shorts for every major activity",
            "Live coverage of District Convention",
            "Bilingual content — English + regional language",
        ],
        "gallery": [
            "Instagram", "Facebook", "LinkedIn",
            "YouTube Reel", "Convention Live", "Hashtag Campaign",
        ],
    },
    {
        "kind": "content",
        "page_label": "PAGE 6 · TOPIC 6",
        "title": "Achievements During the Year",
        "subtitle": "Impact · Growth · Recognition",
        "bullets": [
            "100% MMR compliance for 8 consecutive months",
            "50+ club websites launched under district umbrella",
            "₹50L+ donations processed online — full audit trail",
            "5,000+ service hours digitally logged on MyLion",
            "Recognised at MD 3232 IT Forum",
            "District benefits: transparency · faster comms · trust",
        ],
        "gallery": [
            "Award Moment", "Convention Stage", "Certificate",
            "Metrics Dashboard", "Milestone Post", "Press Feature",
        ],
    },
]


def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_borders(cell, color: str = "003F87", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_width(cell, cm: float) -> None:
    cell.width = Cm(cm)


def _add_run(paragraph, text: str, *, bold=False, italic=False, size=11,
             color: RGBColor | None = None, font: str = "Calibri") -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _logo_header(doc: Document, logos: list[str]) -> None:
    n = len(logos)
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    total_w = 25.0  # cm — landscape A4 usable width approx
    cell_w = total_w / n
    for i, label in enumerate(logos):
        cell = table.cell(0, i)
        _set_cell_width(cell, cell_w)
        _shade(cell, "003F87")
        _set_borders(cell, color="FFC72C", size="12")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, label, bold=True, size=14, color=WHITE, font="Calibri")
    # subheader strip
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub, "LIONS INTERNATIONAL · DISTRICT 3232F1", bold=True, size=12, color=LIONS_BLUE)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub2, "A-19 District Information Technology Chairperson", italic=True, size=10, color=DARK_GREY)


def _add_title_block(doc: Document, page_label: str, title: str, subtitle: str) -> None:
    badge = doc.add_paragraph()
    _add_run(badge, f"  {page_label}  ", bold=True, size=10, color=LIONS_BLUE)
    # color the badge background by inserting a shaded run via a 1x1 table looks cleaner — skip for simplicity

    t = doc.add_paragraph()
    _add_run(t, title, bold=True, size=22, color=LIONS_BLUE)

    s = doc.add_paragraph()
    _add_run(s, subtitle, italic=True, size=12, color=DARK_GREY)


def _add_lead(doc: Document, lead: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade(cell, "F4F6FA")
    _set_borders(cell, color="FFC72C", size="18")
    p = cell.paragraphs[0]
    _add_run(p, lead, italic=False, size=11, color=DARK_GREY)


def _add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        _add_run(p, b, size=11, color=DARK_GREY)


def _placeholder_cell(cell, label: str, index_label: str) -> None:
    _shade(cell, "F4F6FA")
    _set_borders(cell, color="003F87", size="10")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p1, f" {index_label} ", bold=True, size=9, color=LIONS_BLUE)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, "[ Insert Photo ]", italic=True, size=10, color=DARK_GREY)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p3, label, bold=True, size=11, color=LIONS_BLUE)


def _add_governor_row(doc: Document) -> None:
    h = doc.add_paragraph()
    _add_run(h, "Leadership · 2025–26", bold=True, size=13, color=LIONS_BLUE)

    # 1 row x 4 cols horizontal strip
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    cell_w = 25.0 / 4
    for col in range(4):
        cell = table.cell(0, col)
        _set_cell_width(cell, cell_w)
        role, scope = GOVERNORS[col]
        _placeholder_cell(cell, role, f"GOV {col + 1}")
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, scope, italic=True, size=9, color=DARK_GREY)
    table.rows[0].height = Cm(4.5)


def _add_chairperson_photo(doc: Document) -> None:
    h = doc.add_paragraph()
    _add_run(h, "District IT Chairperson · 2025–26",
             bold=True, size=13, color=LIONS_BLUE)

    # Centered single-cell table for the chairperson photo
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    # left + right spacers, centre cell holds the photo
    _set_cell_width(table.cell(0, 0), 7.5)
    _set_cell_width(table.cell(0, 1), 10.0)
    _set_cell_width(table.cell(0, 2), 7.5)
    _placeholder_cell(table.cell(0, 1), "IT Chairperson", "CHAIR")
    p = table.cell(0, 1).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "District 3232F1 · A-19", italic=True, size=9, color=DARK_GREY)
    table.rows[0].height = Cm(5.0)


def _add_gallery_2x3(doc: Document, labels: list[str]) -> None:
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    cell_w = 25.0 / 3
    for r in range(2):
        for col in range(3):
            cell = table.cell(r, col)
            _set_cell_width(cell, cell_w)
            idx = r * 3 + col
            _placeholder_cell(cell, labels[idx], f"IMG {idx + 1}")
        table.rows[r].height = Cm(5.0)


def _footer_line(doc: Document, page_num: int, total: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"We Serve — Powered by Technology   ·   {page_num} / {total}",
             bold=True, size=10, color=LIONS_BLUE)


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Landscape A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width, section.page_height = new_w, new_h
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Document-level footer (applies to all sections by default)
    foot = section.footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(foot, "We Serve — Powered by Technology",
             bold=True, size=10, color=LIONS_BLUE)

    total = len(PAGES)
    for i, page in enumerate(PAGES, start=1):
        if i > 1:
            _page_break(doc)

        logos = COVER_LOGOS if page["kind"] == "cover" else ["DIST 3232F1", "LCI"]
        _logo_header(doc, logos)
        _add_title_block(doc, page["page_label"], page["title"], page["subtitle"])

        if page["kind"] == "cover":
            _add_lead(doc, page["lead"])
            _add_governor_row(doc)
            _add_chairperson_photo(doc)
            _add_bullets(doc, page["bullets"])
        else:
            _add_bullets(doc, page["bullets"])
            _add_gallery_2x3(doc, page["gallery"])

        _footer_line(doc, i, total)

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
